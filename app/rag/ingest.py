import hashlib
from collections.abc import Iterator
from pathlib import Path

import docx
import redis
from docx.table import Table as DocxTable
from docx.text.paragraph import Paragraph as DocxParagraph
from langchain_core.document_loaders import BaseLoader
from langchain_core.documents import Document
from langchain_redis import RedisVectorStore
from langchain_text_splitters import MarkdownHeaderTextSplitter, RecursiveCharacterTextSplitter

from app.core.config import get_settings
from app.integrations.redis import (
    build_redis_client,
    drop_chunk_index,
    read_manifest,
    write_manifest,
)
from app.rag.model import CorpusManifest, IngestResult
from app.rag.store import (
    EMBEDDING_DIMENSION,
    EMBEDDING_MODEL_NAME,
    EMBEDDING_MODEL_REVISION,
    build_embeddings,
    build_vector_store,
)
from app.rules.loader import get_rule_catalogue
from app.rules.model import RuleCatalogue

CORPUS_DIR = Path(".docs/sources/en")
RULES_PATH = Path("rules.yaml")

CHUNK_SIZE = 800
CHUNK_OVERLAP = 120
SHORT_SECTION_MERGE_THRESHOLD = 200
HEADER_LEVELS = [("#", "Header 1"), ("##", "Header 2"), ("###", "Header 3")]
INGEST_BATCH_SIZE = 128

_HEADING_STYLE_TO_MARKER = {"Heading 1": "#", "Heading 2": "##", "Heading 3": "###"}
_LIST_STYLE_TO_MARKER = {"List Bullet": "-", "List Number": "1."}


class IngestionError(RuntimeError):
    """Raised when the corpus cannot be safely converted, chunked or cross-checked."""


def _paragraph_to_markdown(paragraph: DocxParagraph) -> str | None:
    text = paragraph.text.strip()
    if not text:
        return None
    style = paragraph.style.name if paragraph.style else ""
    if style == "Title":
        return f"# {text}"
    heading_marker = _HEADING_STYLE_TO_MARKER.get(style)
    if heading_marker:
        return f"{heading_marker} {text}"
    list_marker = _LIST_STYLE_TO_MARKER.get(style)
    if list_marker:
        return f"{list_marker} {text}"
    return text


def _table_to_markdown(table: DocxTable) -> str:
    rows = [[cell.text.strip().replace("\n", " ") for cell in row.cells] for row in table.rows]
    rows = [row for row in rows if any(row)]
    if not rows:
        return ""
    header, *body = rows
    lines = ["| " + " | ".join(header) + " |", "| " + " | ".join("---" for _ in header) + " |"]
    lines.extend("| " + " | ".join(row) + " |" for row in body)
    return "\n".join(lines)


def convert_docx_to_markdown(path: Path) -> tuple[str, str]:
    """Converts one policy `.docx` file to Markdown. Returns (title, markdown)."""
    document = docx.Document(str(path))
    title = path.stem
    blocks: list[str] = []
    for item in document.iter_inner_content():
        if isinstance(item, DocxParagraph):
            if item.style and item.style.name == "Title" and item.text.strip():
                title = item.text.strip()
            markdown_line = _paragraph_to_markdown(item)
            if markdown_line is not None:
                blocks.append(markdown_line)
        elif isinstance(item, DocxTable):
            table_markdown = _table_to_markdown(item)
            if table_markdown:
                blocks.append(table_markdown)
    return title, "\n\n".join(blocks)


class DocxMarkdownLoader(BaseLoader):
    """LangChain loader converting the fictional company's `.docx` policy corpus to Markdown."""

    def __init__(self, corpus_dir: Path = CORPUS_DIR) -> None:
        self._corpus_dir = corpus_dir

    def lazy_load(self) -> Iterator[Document]:
        for path in sorted(self._corpus_dir.glob("*.docx")):
            doc_id = path.name[:2]
            title, markdown = convert_docx_to_markdown(path)
            yield Document(
                page_content=markdown,
                metadata={"doc_id": doc_id, "doc_title": title, "source_path": str(path)},
            )


def _heading_from_metadata(metadata: dict) -> str | None:
    for key in ("Header 3", "Header 2", "Header 1"):
        if key in metadata:
            return metadata[key]
    return None


def _merge_short_sections(
    sections: list[tuple[str | None, str]],
) -> list[tuple[str | None, str]]:
    """Merges sections shorter than the threshold into the following sibling.

    A trailing short section has no following sibling to merge into, so it is kept as its
    own chunk under its own heading rather than absorbed into an unrelated earlier section.
    """
    merged: list[tuple[str | None, str]] = []
    pending_heading: str | None = None
    pending_text = ""
    for heading, content in sections:
        combined = f"{pending_text}\n\n{content}".strip() if pending_text else content
        if len(combined) < SHORT_SECTION_MERGE_THRESHOLD:
            pending_heading, pending_text = heading, combined
            continue
        merged.append((heading, combined))
        pending_heading, pending_text = None, ""
    if pending_text:
        merged.append((pending_heading, pending_text))
    return merged


def _is_table_line(line: str) -> bool:
    return line.strip().startswith("|")


def _split_prose_and_tables(text: str) -> list[tuple[bool, str]]:
    """Splits section text into ordered segments, keeping table blocks atomic."""
    segments: list[tuple[bool, str]] = []
    buffer: list[str] = []
    buffer_is_table = False

    def flush() -> None:
        segment = "\n".join(buffer).strip()
        if segment:
            segments.append((buffer_is_table, segment))
        buffer.clear()

    for line in text.split("\n"):
        line_is_table = _is_table_line(line)
        if buffer and line_is_table != buffer_is_table:
            flush()
        buffer_is_table = line_is_table
        buffer.append(line)
    flush()
    return segments


def _guard_split_segments(
    segments: list[tuple[bool, str]], splitter: RecursiveCharacterTextSplitter
) -> list[str]:
    chunks: list[str] = []
    for is_table, text in segments:
        chunks.append(text) if is_table else chunks.extend(splitter.split_text(text))
    return chunks


def chunk_markdown_document(
    doc_id: str, doc_title: str, source_path: str, markdown: str
) -> list[Document]:
    """Splits one document's Markdown into header-aware, size-guarded chunks."""
    header_splitter = MarkdownHeaderTextSplitter(HEADER_LEVELS, strip_headers=True)
    header_sections = header_splitter.split_text(markdown)

    raw_sections = [
        (_heading_from_metadata(section.metadata), section.page_content)
        for section in header_sections
        if section.page_content.strip()
    ]
    merged_sections = _merge_short_sections(raw_sections)
    char_splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE, chunk_overlap=CHUNK_OVERLAP
    )

    chunks: list[Document] = []
    for heading, section_text in merged_sections:
        for chunk_text in _guard_split_segments(
            _split_prose_and_tables(section_text), char_splitter
        ):
            chunks.append(
                Document(
                    page_content=chunk_text,
                    metadata={
                        "doc_id": doc_id,
                        "doc_title": doc_title,
                        "section": heading,
                        "chunk_index": len(chunks),
                        "source_path": source_path,
                    },
                )
            )
    return chunks


def _resolve_section_id(
    doc_id: str, heading: str | None, rule_catalogue: RuleCatalogue
) -> str | None:
    if heading is None:
        return None
    document = rule_catalogue.documents.get(doc_id)
    if document is None:
        return None
    for section_id, section in document.sections.items():
        if heading in section.headings:
            return section_id
    return None


def attach_rule_metadata(chunks: list[Document], rule_catalogue: RuleCatalogue) -> list[Document]:
    """Attaches `section_id`, `rule_ids` and `categories` to each chunk's metadata in place."""
    rule_ids_by_section: dict[tuple[str, str], list[str]] = {}
    for category_rules in rule_catalogue.categories.values():
        for rule in category_rules.rules:
            if rule.doc_ref is None:
                continue
            doc_id, section_id = rule.doc_ref.split("#", 1)
            rule_ids_by_section.setdefault((doc_id, section_id), []).append(rule.id)

    for chunk in chunks:
        doc_id = chunk.metadata["doc_id"]
        document = rule_catalogue.documents.get(doc_id)
        if document is None:
            raise IngestionError(
                f"Unknown document identifier {doc_id!r}; not declared in rules.yaml"
            )

        heading = chunk.metadata.get("section")
        section_id = _resolve_section_id(doc_id, heading, rule_catalogue)
        chunk.metadata["section_id"] = section_id
        chunk.metadata["rule_ids"] = (
            rule_ids_by_section.get((doc_id, section_id), []) if section_id else []
        )
        chunk.metadata["categories"] = list(document.categories)

    return chunks


def validate_section_anchors_resolve(chunks: list[Document], rule_catalogue: RuleCatalogue) -> None:
    """Rejects rules.yaml section anchors whose heading never appears in the ingested corpus."""
    headings_by_doc: dict[str, set[str]] = {}
    for chunk in chunks:
        heading = chunk.metadata.get("section")
        if heading:
            headings_by_doc.setdefault(chunk.metadata["doc_id"], set()).add(heading)

    errors = [
        f"rules.yaml section '{doc_id}#{section_id}' heading {heading!r} not in the corpus"
        for doc_id, document in rule_catalogue.documents.items()
        for section_id, section in document.sections.items()
        for heading in section.headings
        if heading not in headings_by_doc.get(doc_id, set())
    ]
    if errors:
        raise IngestionError("; ".join(errors))


def load_and_chunk_corpus(
    corpus_dir: Path, rule_catalogue: RuleCatalogue
) -> tuple[list[Document], list[Document]]:
    """Returns (source_documents, chunks) for the corpus, validated against the rule catalogue."""
    source_documents = list(DocxMarkdownLoader(corpus_dir).lazy_load())
    chunks = [
        chunk
        for source in source_documents
        for chunk in chunk_markdown_document(
            source.metadata["doc_id"],
            source.metadata["doc_title"],
            source.metadata["source_path"],
            source.page_content,
        )
    ]
    attach_rule_metadata(chunks, rule_catalogue)
    validate_section_anchors_resolve(chunks, rule_catalogue)
    return source_documents, chunks


def compute_corpus_hash(corpus_dir: Path, rules_path: Path) -> str:
    hasher = hashlib.sha256()
    for path in sorted(corpus_dir.glob("*.docx")):
        hasher.update(path.name.encode())
        hasher.update(path.read_bytes())
    hasher.update(rules_path.read_bytes())
    return hasher.hexdigest()


def build_manifest(
    corpus_dir: Path,
    rules_path: Path,
    embedding_model: str,
    embedding_revision: str,
    dimension: int,
) -> CorpusManifest:
    return CorpusManifest(
        corpus_hash=compute_corpus_hash(corpus_dir, rules_path),
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
        short_section_merge_threshold=SHORT_SECTION_MERGE_THRESHOLD,
        embedding_model=embedding_model,
        embedding_revision=embedding_revision,
        dimension=dimension,
    )


def _count_categories(chunks: list[Document]) -> dict[str, int]:
    counts: dict[str, int] = {}
    for chunk in chunks:
        for category in chunk.metadata["categories"]:
            counts[category] = counts.get(category, 0) + 1
    return counts


def _upsert_chunks(vector_store: RedisVectorStore, chunks: list[Document]) -> None:
    for start in range(0, len(chunks), INGEST_BATCH_SIZE):
        batch = chunks[start : start + INGEST_BATCH_SIZE]
        vector_store.add_texts(
            texts=[chunk.page_content for chunk in batch],
            metadatas=[chunk.metadata for chunk in batch],
            keys=[f"{chunk.metadata['doc_id']}:{chunk.metadata['chunk_index']}" for chunk in batch],
        )


def run_ingest(
    redis_client: redis.Redis,
    vector_store: RedisVectorStore,
    corpus_dir: Path = CORPUS_DIR,
    rules_path: Path = RULES_PATH,
    rule_catalogue: RuleCatalogue | None = None,
) -> IngestResult:
    """Ingests the corpus into Redis, skipping the embed/upsert step when the manifest matches."""
    rule_catalogue = rule_catalogue or get_rule_catalogue()
    _, chunks = load_and_chunk_corpus(corpus_dir, rule_catalogue)

    manifest = build_manifest(
        corpus_dir, rules_path, EMBEDDING_MODEL_NAME, EMBEDDING_MODEL_REVISION, EMBEDDING_DIMENSION
    )
    existing_manifest = read_manifest(redis_client)

    if existing_manifest == manifest:
        action = "reused"
    else:
        action = "rebuilt" if existing_manifest is not None else "built"
        if existing_manifest is not None:
            drop_chunk_index(redis_client)
        _upsert_chunks(vector_store, chunks)
        write_manifest(redis_client, manifest)

    return IngestResult(
        action=action, chunk_count=len(chunks), category_counts=_count_categories(chunks)
    )


if __name__ == "__main__":
    settings = get_settings()
    redis_url = settings.redis_url
    result = run_ingest(
        build_redis_client(redis_url),
        build_vector_store(redis_url, build_embeddings()),
        rule_catalogue=get_rule_catalogue(),
    )
    print(result.model_dump_json(indent=2))
