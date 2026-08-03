import docx
import pytest
from docx.opc.exceptions import PackageNotFoundError

from app.rag.ingest.docx_converter import DocxToMarkdownConverter


def test_convert_returns_empty_markdown_for_a_document_with_no_content(tmp_path):
    # Arrange
    path = tmp_path / "01_empty.docx"
    docx.Document().save(str(path))

    # Act
    title, markdown = DocxToMarkdownConverter().convert(path)

    # Assert
    assert title == "01_empty"
    assert markdown == ""


def test_table_cell_newlines_are_collapsed_to_spaces(tmp_path):
    # Arrange
    path = tmp_path / "01_sample.docx"
    document = docx.Document()
    table = document.add_table(rows=2, cols=1)
    table.cell(0, 0).text = "Header"
    table.cell(1, 0).text = "Line one\nLine two"
    document.save(str(path))

    # Act
    _, markdown = DocxToMarkdownConverter().convert(path)

    # Assert
    assert markdown == "| Header |\n| --- |\n| Line one Line two |"


def test_table_with_only_a_header_row_renders_without_body_rows(tmp_path):
    # Arrange
    path = tmp_path / "01_sample.docx"
    document = docx.Document()
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Category"
    table.cell(0, 1).text = "Limit"
    document.save(str(path))

    # Act
    _, markdown = DocxToMarkdownConverter().convert(path)

    # Assert
    assert markdown == "| Category | Limit |\n| --- | --- |"


def test_convert_raises_for_a_malformed_docx_file(tmp_path):
    # Arrange
    path = tmp_path / "01_broken.docx"
    path.write_text("this is not a real docx file")

    # Act / Assert
    converter = DocxToMarkdownConverter()
    with pytest.raises(PackageNotFoundError):
        converter.convert(path)
