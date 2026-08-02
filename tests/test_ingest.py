import docx
import pytest
from langchain_core.documents import Document

from app.rag.build_info import IndexBuildInfoBuilder
from app.rag.chunker import CHUNK_SIZE, MarkdownChunker
from app.rag.docx_converter import DocxToMarkdownConverter
from app.rag.errors import IngestionError
from app.rag.ingest import CorpusIngestor
from app.rag.rule_metadata import RuleMetadataResolver
from app.rules.loader import load_rule_catalogue
from app.rules.model import RuleCatalogue

RULES_FIXTURE = {
    "version": 1,
    "currency": "HUF",
    "fx_rates_fixed": {"EUR": 400},
    "documents": {
        "01": {
            "categories": ["meal"],
            "sections": {"limit": {"headings": ["4. Business meals"]}},
        }
    },
    "submission": {"deadline_days": 30, "approval_tiers": [{"max_huf": None, "approver": "x"}]},
    "categories": {
        "meal": {
            "rules": [{"id": "R-MEAL-01", "doc_ref": "01#limit"}],
            "required_documents": [],
        }
    },
}


def _build_docx(path, add_title=True):
    document = docx.Document()
    if add_title:
        document.add_paragraph("Sample Policy", style="Title")
    document.add_paragraph("1. Purpose and scope", style="Heading 1")
    document.add_paragraph("This section explains the purpose.")
    document.add_paragraph("First bullet", style="List Bullet")
    document.add_paragraph("Second bullet", style="List Bullet")
    document.add_paragraph("4. Business meals", style="Heading 1")
    document.add_paragraph("Meals are capped per person per occasion.")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Category"
    table.cell(0, 1).text = "Limit"
    table.cell(1, 0).text = "Meal"
    table.cell(1, 1).text = "HUF 15,000"
    document.save(str(path))
    return path


def test_convert_docx_to_markdown_maps_headings_lists_tables_and_title(tmp_path):
    path = _build_docx(tmp_path / "01_sample.docx")

    title, markdown = DocxToMarkdownConverter().convert(path)

    assert title == "Sample Policy"
    assert "# Sample Policy" in markdown
    assert "# 1. Purpose and scope" in markdown
    assert "- First bullet" in markdown
    assert "- Second bullet" in markdown
    assert "# 4. Business meals" in markdown
    assert "| Category | Limit |" in markdown
    assert "| Meal | HUF 15,000 |" in markdown


def test_convert_docx_to_markdown_falls_back_to_filename_without_title_style(tmp_path):
    path = _build_docx(tmp_path / "01_sample.docx", add_title=False)

    title, _ = DocxToMarkdownConverter().convert(path)

    assert title == "01_sample"


def test_table_is_kept_whole_even_when_it_exceeds_chunk_size():
    long_row = " | ".join(f"cell{i}" * 60 for i in range(4))
    markdown = f"# 1. Big table\n\n| a | b | c | d |\n| --- | --- | --- | --- |\n| {long_row} |"

    chunks = MarkdownChunker().chunk("01", "Doc", "source.docx", markdown)

    table_chunks = [c for c in chunks if "| a |" in c.page_content]
    assert len(table_chunks) == 1
    assert table_chunks[0].page_content.startswith("1. Big table\n\n| a |")
    assert len(table_chunks[0].page_content) > CHUNK_SIZE


def test_short_section_is_merged_into_the_following_sibling():
    markdown = "# 1. Tiny\n\nOK.\n\n# 2. Real section\n\n" + (
        "This is a normal-length paragraph about the real policy. " * 6
    )

    chunks = MarkdownChunker().chunk("01", "Doc", "source.docx", markdown)

    assert all(c.metadata["section"] != "1. Tiny" for c in chunks)
    assert any("OK." in c.page_content for c in chunks)
    assert any(c.metadata["section"] == "2. Real section" for c in chunks)


def test_long_prose_section_is_split_into_multiple_overlapping_chunks():
    paragraph = "This sentence repeats to build a long section of prose text. "
    markdown = "# 1. Long section\n\n" + paragraph * 30

    chunks = MarkdownChunker().chunk("01", "Doc", "source.docx", markdown)

    assert len(chunks) > 1
    assert all(c.metadata["section"] == "1. Long section" for c in chunks)
    assert all(c.page_content.startswith("1. Long section\n\n") for c in chunks)
    heading_overhead = len("1. Long section\n\n")
    assert all(len(c.page_content) <= CHUNK_SIZE + heading_overhead for c in chunks)


def test_chunk_index_is_sequential_per_document():
    markdown = "# 1. A\n\nFirst.\n\n# 2. B\n\nSecond."

    chunks = MarkdownChunker().chunk("01", "Doc", "source.docx", markdown)

    assert [c.metadata["chunk_index"] for c in chunks] == list(range(len(chunks)))


def test_rule_metadata_resolver_rejects_unknown_document_id():
    catalogue = RuleCatalogue.model_validate(RULES_FIXTURE)
    chunk = Document(page_content="x", metadata={"doc_id": "99", "section": None})
    resolver = RuleMetadataResolver(catalogue)

    with pytest.raises(IngestionError, match="Unknown document identifier"):
        resolver.attach([chunk])


def test_rule_metadata_resolver_assigns_categories_and_rule_ids():
    catalogue = RuleCatalogue.model_validate(RULES_FIXTURE)
    chunk = Document(page_content="x", metadata={"doc_id": "01", "section": "4. Business meals"})

    RuleMetadataResolver(catalogue).attach([chunk])

    assert chunk.metadata["categories"] == ["meal"]
    assert chunk.metadata["section_id"] == "limit"
    assert chunk.metadata["rule_ids"] == ["R-MEAL-01"]


def test_validate_anchors_resolve_rejects_missing_heading():
    catalogue = RuleCatalogue.model_validate(RULES_FIXTURE)
    chunk = Document(page_content="x", metadata={"doc_id": "01", "section": "Unrelated heading"})
    resolver = RuleMetadataResolver(catalogue)

    with pytest.raises(IngestionError, match="not in the corpus"):
        resolver.validate_anchors_resolve([chunk])


def test_validate_anchors_resolve_passes_when_heading_present():
    catalogue = RuleCatalogue.model_validate(RULES_FIXTURE)
    chunk = Document(page_content="x", metadata={"doc_id": "01", "section": "4. Business meals"})

    RuleMetadataResolver(catalogue).validate_anchors_resolve([chunk])


def test_validate_categories_reachable_passes_when_every_category_has_a_chunk():
    catalogue = RuleCatalogue.model_validate(RULES_FIXTURE)
    chunk = Document(
        page_content="x",
        metadata={"doc_id": "01", "section": "4. Business meals", "categories": ["meal"]},
    )

    resolver = RuleMetadataResolver(catalogue)
    resolver.attach([chunk])
    resolver.validate_categories_reachable([chunk])


def test_validate_categories_reachable_rejects_a_category_with_no_chunk():
    fixture = {
        **RULES_FIXTURE,
        "categories": {
            **RULES_FIXTURE["categories"],
            "equipment": {"rules": [], "required_documents": []},
        },
    }
    catalogue = RuleCatalogue.model_validate(fixture)
    chunk = Document(
        page_content="x",
        metadata={"doc_id": "01", "section": "4. Business meals", "categories": ["meal"]},
    )
    resolver = RuleMetadataResolver(catalogue)
    resolver.attach([chunk])

    with pytest.raises(IngestionError, match="equipment"):
        resolver.validate_categories_reachable([chunk])


def test_cross_document_rule_adds_its_own_category_to_the_referenced_section():
    catalogue = load_rule_catalogue()
    chunk = Document(
        page_content="Traffic and parking fines are not reimbursable.",
        metadata={"doc_id": "01", "section": "6. Non-reimbursable items"},
    )

    RuleMetadataResolver(catalogue).attach([chunk])

    assert "travel" in chunk.metadata["categories"]
    assert "R-TRAVEL-04" in chunk.metadata["rule_ids"]


def test_validate_categories_reachable_rejects_missing_rule_evidence():
    catalogue = RuleCatalogue.model_validate(RULES_FIXTURE)
    chunk = Document(page_content="x", metadata={"doc_id": "01", "section": "Other"})
    resolver = RuleMetadataResolver(catalogue)
    resolver.attach([chunk])

    with pytest.raises(IngestionError, match="meal:R-MEAL-01"):
        resolver.validate_categories_reachable([chunk])


def test_compute_corpus_hash_changes_when_a_document_changes(tmp_path):
    corpus_dir = tmp_path / "corpus"
    corpus_dir.mkdir()
    rules_path = tmp_path / "rules.yaml"
    rules_path.write_text("version: 1\n")
    _build_docx(corpus_dir / "01_sample.docx")
    builder = IndexBuildInfoBuilder(corpus_dir, rules_path)

    first_hash = builder.compute_hash()
    second_hash = builder.compute_hash()
    assert first_hash == second_hash

    _build_docx(corpus_dir / "02_other.docx")
    third_hash = builder.compute_hash()
    assert third_hash != first_hash


def test_compute_corpus_hash_changes_when_rules_change(tmp_path):
    corpus_dir = tmp_path / "corpus"
    corpus_dir.mkdir()
    _build_docx(corpus_dir / "01_sample.docx")
    rules_path = tmp_path / "rules.yaml"
    builder = IndexBuildInfoBuilder(corpus_dir, rules_path)

    rules_path.write_text("version: 1\n")
    first_hash = builder.compute_hash()

    rules_path.write_text("version: 2\n")
    second_hash = builder.compute_hash()

    assert first_hash != second_hash


def test_load_and_chunk_end_to_end(tmp_path):
    corpus_dir = tmp_path / "corpus"
    corpus_dir.mkdir()
    _build_docx(corpus_dir / "01_sample.docx")
    catalogue = RuleCatalogue.model_validate(RULES_FIXTURE)

    sources, chunks = CorpusIngestor(corpus_dir).load_and_chunk(catalogue)

    assert len(sources) == 1
    assert sources[0].metadata["doc_id"] == "01"
    assert any(c.metadata["rule_ids"] == ["R-MEAL-01"] for c in chunks)


def test_build_info_builder_reflects_chunking_and_embedding_settings(tmp_path):
    corpus_dir = tmp_path / "corpus"
    corpus_dir.mkdir()
    _build_docx(corpus_dir / "01_sample.docx")
    rules_path = tmp_path / "rules.yaml"
    rules_path.write_text("version: 1\n")

    build_info = IndexBuildInfoBuilder(corpus_dir, rules_path).build("model-x", "rev-1", 384)

    assert build_info.chunk_size == CHUNK_SIZE
    assert build_info.embedding_model == "model-x"
    assert build_info.embedding_revision == "rev-1"
    assert build_info.dimension == 384
